"""Document parsing: dispatch by mime type / extension to a text extractor.

Each parser returns plain text. PDF extraction inserts form-feed-marked page
breaks so downstream metadata can record approximate page numbers if desired.
"""
from __future__ import annotations

import io
from datetime import date, datetime


class UnsupportedDocumentError(ValueError):
    """Raised when no parser matches the given content type."""


def parse_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = [(page.extract_text() or "").strip() for page in reader.pages]
    return "\n\n".join(p for p in pages if p)


def parse_docx(data: bytes) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(data))
    paragraphs = [p.text.strip() for p in doc.paragraphs]
    return "\n\n".join(p for p in paragraphs if p)


def parse_text(data: bytes) -> str:
    return data.decode("utf-8", errors="replace").strip()


# Resolution order: explicit mime type first, then filename extension.
_MIME_DISPATCH = {
    "application/pdf": parse_pdf,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": parse_docx,
    "text/plain": parse_text,
    "text/markdown": parse_text,
}

_EXT_DISPATCH = {
    ".pdf": parse_pdf,
    ".docx": parse_docx,
    ".txt": parse_text,
    ".md": parse_text,
    ".markdown": parse_text,
}


def parse_document(data: bytes, *, mime_type: str = "", filename: str = "") -> str:
    """Extract plain text from raw bytes using mime type or extension."""
    parser = _MIME_DISPATCH.get(mime_type.split(";")[0].strip().lower())
    if parser is None and filename:
        ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        parser = _EXT_DISPATCH.get(ext)
    if parser is None:
        raise UnsupportedDocumentError(
            f"No parser for mime_type={mime_type!r} filename={filename!r}"
        )
    return parser(data)


def _coerce_date(value) -> date | None:
    if isinstance(value, datetime):
        return value.date()
    if isinstance(value, date):
        return value
    return None


def extract_metadata(data: bytes, *, mime_type: str = "", filename: str = "") -> dict:
    """Best-effort document-level metadata (author, date) from file properties.

    Never raises: extraction is advisory, so failures degrade to empty values.
    Returned keys are only present when a value was found.
    """
    meta: dict = {}
    mt = mime_type.split(";")[0].strip().lower()
    ext = ("." + filename.rsplit(".", 1)[-1].lower()) if "." in filename else ""

    try:
        if mt == "application/pdf" or ext == ".pdf":
            from pypdf import PdfReader

            info = PdfReader(io.BytesIO(data)).metadata
            if info:
                if info.author:
                    meta["author"] = str(info.author)
                created = _coerce_date(getattr(info, "creation_date", None))
                if created:
                    meta["doc_date"] = created
        elif mt.endswith("wordprocessingml.document") or ext == ".docx":
            from docx import Document as DocxDocument

            props = DocxDocument(io.BytesIO(data)).core_properties
            if props.author:
                meta["author"] = str(props.author)
            created = _coerce_date(props.created)
            if created:
                meta["doc_date"] = created
    except Exception:  # noqa: BLE001 — metadata extraction is best-effort
        return meta
    return meta
